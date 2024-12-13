import streamlit as st
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def replace_placeholders(doc, placeholders):
    """Replace placeholders in a document while maintaining proper alignment."""
    # Keywords to detect left-side content
    left_side_keywords = [
        "BILL TO", "Mobile No", "Address", "Email", "Project Name", "Company Name"
    ]

    # Iterate through all paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                inline = para.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(key, value)
                # Force left alignment for specific placeholders
                if any(keyword in para.text for keyword in left_side_keywords):
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.left_indent = None  # Reset any indent
                    para.paragraph_format.first_line_indent = None  # Reset first-line indent

    # Iterate through all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in placeholders.items():
                        if key in para.text:
                            inline = para.runs
                            for i in range(len(inline)):
                                if key in inline[i].text:
                                    inline[i].text = inline[i].text.replace(key, value)
                            # Force left alignment for specific placeholders in tables
                            if any(keyword in para.text for keyword in left_side_keywords):
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                para.paragraph_format.left_indent = None
                                para.paragraph_format.first_line_indent = None

    return doc

    return doc


def format_percentage(value):
    """Format percentage without decimals."""
    return f"{int(value)}%"


def edit_invoice_template(template_name, output_path, placeholders):
    """Edit an invoice template and save the result."""
    try:
        doc = Document(template_name)
        replace_placeholders(doc, placeholders)
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing invoice template: {e}")

def format_price(price, currency):
    """Format price to display correctly with the currency."""
    if price.is_integer():
        formatted_price = f"{int(price)}"
    else:
        formatted_price = f"{price:.2f}"
    if currency == "USD":
        return f"{formatted_price} USD"
    elif currency == "Rupees":
        return f"Rs. {formatted_price}"
    return formatted_price


def generate_invoice():
    """Streamlit app for generating invoices."""
    st.title("Invoice Generator")

    # Reordered Fields
    region = st.selectbox("Region", ["ROW", "India"])
    client_name = st.text_input("Client Name")
    company_name = st.text_input("Company Name")
    contact = st.text_input("Contact Number")
    address = st.text_area("Address")
    project_name = st.text_input("Project Name")
    email = st.text_input("Email")
    service = st.text_input("Service")
    currency = st.selectbox("Currency", ["USD", "Rupees"])
    total_amount = st.number_input("Total Amount", min_value=0.0, format="%.2f")
    payment_option = st.selectbox("Payment Option", ["One Part", "Two Parts", "Three Parts"])
    invoice_date = st.date_input("Invoice Date", value=datetime.today())

    # Service Description for "One Part"
    service_description = ""
    if payment_option == "One Part":
        service_description = st.text_area("Service Description (Optional)")

    # Collect additional inputs for multi-part payments
    if payment_option == "Two Parts":
        p1_percentage = st.number_input("Percentage for First Installment", min_value=0.0, max_value=100.0)
        p2_percentage = 100 - p1_percentage
    elif payment_option == "Three Parts":
        p1_percentage = st.number_input("Percentage for First Installment", min_value=0.0, max_value=100.0)
        max_p2 = 100 - p1_percentage
        p2_percentage = st.number_input("Percentage for Second Installment", min_value=0.0, max_value=max_p2)
        p3_percentage = 100 - (p1_percentage + p2_percentage)

    # Calculate payment amounts
    # Calculate payment amounts
    if payment_option == "Two Parts":
       p1_percentage = round(p1_percentage)
       p2_percentage = 100 - p1_percentage
       price = round(total_amount * (p1_percentage / 100))
       price2 = total_amount - price
       
    elif payment_option == "Three Parts":
       p1_percentage = round(p1_percentage)
       p2_percentage = round(p2_percentage)
       p3_percentage = 100 - (p1_percentage + p2_percentage)
       price = round(total_amount * (p1_percentage / 100))
       price2 = round(total_amount * (p2_percentage / 100))
       price3 = total_amount - (price + price2)



    # Generate placeholders
    formatted_date = invoice_date.strftime("%d/%m/%Y")
    placeholders = {
        "<< Client Name >>": client_name,
        "<<Company Name>>": company_name,
        "<<Client Contact>>": contact,
        "<<Address>>": address,
        "<<Client Email>>": email,
        "<<Project Name>>": project_name,
        "<<Service>>": service,
        "<<Price>>": format_price(total_amount, currency),
        "<< Date >>": formatted_date,
    }

    if service_description:
        placeholders["<<Service Description>>"] = service_description

    if payment_option == "Two Parts":
        placeholders.update({
            "<<P1>>": format_percentage(p1_percentage),
            "<<Price>>": format_price(price, currency),
            "<<P2>>": format_percentage(p2_percentage),
            "<<Price2>>": format_price(price2, currency),
       })
    elif payment_option == "Three Parts":
        placeholders.update({
            "<<P1>>": format_percentage(p1_percentage),
            "<<Price>>": format_price(price, currency),
            "<<P2>>": format_percentage(p2_percentage),
            "<<Price2>>": format_price(price2, currency),
            "<<P3>>": format_percentage(p3_percentage),
            "<<Price3>>": format_price(price3, currency),
        })


    # Select template based on region and payment option
    if payment_option == "One Part" and not service_description.strip():
        template_name = {
            "ROW": "One Part Payment ROW no service.docx",
            "India": "One Part Payment INDIA no service.docx",
        }[region]
    else:
        template_name = {
            "One Part": {
                "ROW": "One Part Payment ROW.docx",
                "India": "One Part Payment INDIA.docx",
            },
            "Two Parts": {
                "ROW": "Two Parts Payment ROW.docx",
                "India": "Two Parts Payment INDIA.docx",
            },
            "Three Parts": {
                "ROW": "Three Parts Payment ROW.docx",
                "India": "Three Parts Payment INDIA.docx",
            },
        }[payment_option][region]

    if st.button("Generate Invoice"):
        formatted_date_filename = invoice_date.strftime("%d %b %Y")  # Adjust filename date format
        output_path = f"Invoice - {client_name} {formatted_date_filename}.docx"
        try:
            edit_invoice_template(template_name, output_path, placeholders)
            st.success("Invoice generated successfully!")
            with open(output_path, "rb") as file:
                st.download_button(
                    label="Download Invoice",
                    data=file,
                    file_name=output_path,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"An error occurred: {e}")

# Execute the app
generate_invoice()
